import os
from flask import Flask, render_template, request, redirect, url_for, jsonify, send_file
from werkzeug.utils import secure_filename
# ADDED: Import the json module
import json
from parser import extract_text_from_pdf, extract_text_from_docx, parse_resume
# Import database functions (now using Firestore)
from database import insert_resume_data, get_all_resumes, delete_resume_data, get_resume_by_id 
import uuid # Import uuid for unique filenames
import io
import re # Import re for filename cleaning

# For PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER

# For DOCX generation
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

# Configuration for file uploads
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Create the uploads directory if it doesn't exist
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Note: Firebase initialization now happens directly in database.py when it's imported.
# No explicit init_db() call needed here.

def allowed_file(filename):
    """
    Checks if the uploaded file has an allowed extension.
    """
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    """
    Renders the main upload page.
    """
    return render_template('index.html')

@app.route('/parse', methods=['POST'])
def parse():
    """
    Handles resume file upload, parsing, and displays the auto-filled form.
    """
    if 'resume' not in request.files:
        print("Error: No file part in request.") # Debugging print
        return jsonify({'error': 'No file part'}), 400

    file = request.files['resume']

    if file.filename == '':
        print("Error: No selected file.") # Debugging print
        return jsonify({'error': 'No selected file'}), 400

    if file and allowed_file(file.filename):
        original_filename = secure_filename(file.filename)
        # Generate a unique filename to prevent caching issues and overwrites
        unique_filename = str(uuid.uuid4()) + "_" + original_filename
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        
        try:
            file.save(filepath)
            print(f"DEBUG: File saved to: {filepath}") # Debugging print

            extracted_text = ""
            if original_filename.endswith('.pdf'):
                extracted_text = extract_text_from_pdf(filepath)
            elif original_filename.endswith('.docx'):
                extracted_text = extract_text_from_docx(filepath)

            if not extracted_text:
                print(f"ERROR: Could not extract text from {unique_filename}. File might be empty or corrupted.") # Debugging print
                return jsonify({'error': 'Could not extract text from the document. The file might be empty or corrupted.'}), 500

            print(f"DEBUG: Extracted text length from {unique_filename}: {len(extracted_text)} characters.") # Debugging print

            parsed_data = parse_resume(extracted_text)
            
            # Store the original filename in the parsed_data dictionary
            parsed_data['original_filename'] = original_filename

            # Debugging: Print the parsed data before rendering the template
            print("\n--- DEBUG: Parsed Data before rendering parsed_form.html ---")
            print(json.dumps(parsed_data, indent=2)) # json.dumps requires 'json' module
            print("---------------------------------------------------\n")

            # Render a new template to display the parsed data in an editable form
            return render_template('parsed_form.html', data=parsed_data)

        except Exception as e:
            print(f"ERROR: An error occurred during file processing: {str(e)}") # Debugging print
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500
        finally:
            # Ensure the uploaded file is removed after processing
            if os.path.exists(filepath):
                os.remove(filepath)
                print(f"DEBUG: Cleaned up: {filepath}") # Debugging print
    else:
        print(f"ERROR: Unsupported file type for {original_filename if 'original_filename' in locals() else 'unknown file'}.") # Debugging print
        return jsonify({'error': 'Unsupported file type. Please upload a PDF or DOCX file.'}), 400

@app.route('/submit_form', methods=['POST'])
def submit_form():
    """
    Handles the submission of the edited form data and saves it to the database.
    """
    # Get form data from request.form
    name = request.form.get('name', '')
    email = request.form.get('email', '')
    phone = request.form.get('phone', '')

    # Skills are submitted as a comma-separated string, convert to list
    skills_str = request.form.get('skills', '')
    skills = [s.strip() for s in skills_str.split(',') if s.strip()]

    education = []
    i = 0
    while True:
        institution = request.form.get(f'education_{i}_institution')
        degree = request.form.get(f'education_{i}_degree')
        passing_year = request.form.get(f'education_{i}_passing_year')
        marks_percentage_cgpa = request.form.get(f'education_{i}_marks_percentage_cgpa')
        
        # Break if no education field for this index is found
        if institution is None and degree is None and passing_year is None and marks_percentage_cgpa is None:
            break
        
        # Only add if at least one field is present for the current education entry
        if institution or degree or passing_year or marks_percentage_cgpa:
            education.append({
                'institution': institution if institution else '',
                'degree': degree if degree else '',
                'passing_year': passing_year if passing_year else '',
                'marks_percentage_cgpa': marks_percentage_cgpa if marks_percentage_cgpa else ''
            })
        i += 1

    work_experience = []
    j = 0
    while True:
        company = request.form.get(f'work_experience_{j}_company')
        title = request.form.get(f'work_experience_{j}_title')
        dates = request.form.get(f'work_experience_{j}_dates')
        description = request.form.get(f'work_experience_{j}_description')
        
        # Break if no work experience field for this index is found
        if company is None and title is None and dates is None and description is None:
            break
        
        # Only add if at least one field is present for the current experience entry
        if company or title or dates or description:
            work_experience.append({
                'company': company if company else '',
                'title': title if title else '',
                'dates': dates if dates else '',
                'description': description if description else ''
            })
        j += 1

    final_data = {
        'name': name,
        'email': email,
        'phone': phone,
        'skills': skills,
        'education': education,
        'work_experience': work_experience
    }

    # Get the original filename passed from the hidden field
    original_filename = request.form.get('original_filename')

    # Save the data to the database (now Firestore)
    try:
        inserted_id = insert_resume_data(final_data, original_filename)
        print(f"DEBUG: Data saved to Firestore with ID: {inserted_id}")
    except Exception as e:
        print(f"ERROR: Error saving data to Firestore: {e}")
        return jsonify({'error': f'Failed to save data to Firestore: {str(e)}'}), 500

    # Redirect to the view_resumes page after successful submission
    return redirect(url_for('view_resumes', db_save_status="success"))

@app.route('/view_resumes')
def view_resumes():
    """
    Fetches and displays all saved resumes from the database (now Firestore).
    """
    all_resumes = get_all_resumes()
    db_save_status = request.args.get('db_save_status') # Get status from redirect
    return render_template('view_resumes.html', resumes=all_resumes, db_save_status=db_save_status)

@app.route('/view_resume_detail/<string:resume_id>') # Changed to string for Firestore ID
def view_resume_detail(resume_id):
    """
    Fetches and displays a single resume's details from the database (now Firestore).
    """
    resume = get_resume_by_id(resume_id)
    
    if resume:
        return render_template('result.html', data=resume, db_save_status="view_mode")
    else:
        return "Resume not found", 404

@app.route('/download_resume/<string:resume_id>/<file_type>') # Changed to string for Firestore ID
def download_resume(resume_id, file_type):
    """
    Generates and serves a resume file (PDF or DOCX) based on database data (now Firestore).
    """
    resume_data = get_resume_by_id(resume_id)
    if not resume_data:
        return "Resume not found", 404

    buffer = io.BytesIO()
    filename_base = resume_data.get('name', 'resume').replace(' ', '_') # Use name for filename, secure_filename not needed for base
    # Ensure filename_base is safe for file systems (remove invalid chars)
    filename_base = re.sub(r'[^\w\s.-]', '', filename_base).strip()
    filename_base = re.sub(r'\s+', '_', filename_base)


    if file_type == 'pdf':
        generate_pdf_from_data(resume_data, buffer)
        buffer.seek(0)
        return send_file(buffer, mimetype='application/pdf', as_attachment=True, download_name=f'{filename_base}.pdf')
    elif file_type == 'docx':
        generate_docx_from_data(resume_data, buffer)
        buffer.seek(0)
        return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=f'{filename_base}.docx')
    else:
        return "Invalid file type", 400

@app.route('/delete_resume/<string:resume_id>', methods=['POST']) # Changed to string for Firestore ID
def delete_resume(resume_id):
    """
    Deletes a resume record from the database (now Firestore).
    """
    delete_resume_data(resume_id)
    return redirect(url_for('view_resumes', db_save_status="deleted"))


# Helper functions for resume generation (no changes needed here for Firestore switch)
def generate_pdf_from_data(data, buffer):
    """Generates a PDF resume from parsed data and writes to a buffer."""
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Modify existing 'Normal' and 'Bullet' styles
    styles['Normal'].fontSize = 10
    styles['Normal'].leading = 12
    styles['Normal'].spaceAfter = 6

    styles['Bullet'].fontSize = 10
    styles['Bullet'].leading = 12
    styles['Bullet'].spaceAfter = 2
    styles['Bullet'].leftIndent = 0.2 * inch
    styles['Bullet'].bulletIndent = 0.1 * inch
    styles['Bullet'].bulletText = '•'

    # Add custom styles with unique names
    styles.add(ParagraphStyle(name='TitleStyle', fontSize=24, leading=28, alignment=TA_CENTER, spaceAfter=12, fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='ContactStyle', fontSize=10, leading=12, alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle(name='SectionHeading', fontSize=16, leading=18, spaceAfter=8, fontName='Helvetica-Bold', textColor='#4f46e5')) # Indigo color
    styles.add(ParagraphStyle(name='SubHeading', fontSize=12, leading=14, spaceAfter=4, fontName='Helvetica-Bold'))
    
    story = []

    # Personal Details
    story.append(Paragraph(data.get('name', 'Unnamed Resume'), styles['TitleStyle']))
    contact_info = []
    if data.get('email'):
        contact_info.append(data['email'])
    if data.get('phone'):
        contact_info.append(data['phone'])
    if contact_info:
        story.append(Paragraph(" | ".join(contact_info), styles['ContactStyle']))
    story.append(Spacer(1, 0.2 * inch))

    # Skills
    if data.get('skills'):
        story.append(Paragraph("Skills", styles['SectionHeading']))
        story.append(Paragraph(", ".join(data['skills']), styles['Normal']))
        story.append(Spacer(1, 0.1 * inch))

    # Education
    if data.get('education'):
        story.append(Paragraph("Education", styles['SectionHeading']))
        for edu in data['education']:
            edu_line = f"<b>{edu.get('degree', '')}</b>"
            if edu.get('institution'):
                edu_line += f" - {edu.get('institution')}"
            if edu.get('passing_year'):
                edu_line += f" ({edu.get('passing_year')})"
            story.append(Paragraph(edu_line, styles['SubHeading']))
            if edu.get('marks_percentage_cgpa'):
                story.append(Paragraph(f"Marks/CGPA: {edu.get('marks_percentage_cgpa')}", styles['Normal']))
            story.append(Spacer(1, 0.05 * inch))
        story.append(Spacer(1, 0.1 * inch))

    # Work Experience
    if data.get('work_experience'):
        story.append(Paragraph("Work Experience", styles['SectionHeading']))
        for exp in data['work_experience']:
            exp_line = f"<b>{exp.get('title', '')}</b>"
            if exp.get('company'):
                exp_line += f" at {exp.get('company')}"
            if exp.get('dates'):
                exp_line += f" ({exp.get('dates')})"
            story.append(Paragraph(exp_line, styles['SubHeading']))
            if exp.get('description'):
                # Split description into lines and add as bullet points
                desc_lines = [line.strip() for line in exp['description'].split('\n') if line.strip()]
                for desc_line in desc_lines:
                    story.append(Paragraph(desc_line, styles['Bullet']))
            story.append(Spacer(1, 0.05 * inch))
        story.append(Spacer(1, 0.1 * inch))

    doc.build(story)

def generate_docx_from_data(data, buffer):
    """Generates a DOCX resume from parsed data and writes to a buffer."""
    document = DocxDocument()
    
    # Add Name
    name_para = document.add_paragraph()
    name_run = name_para.add_run(data.get('name', 'Unnamed Resume'))
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Contact Info
    contact_para = document.add_paragraph()
    contact_info = []
    if data.get('email'):
        contact_info.append(data['email'])
    if data.get('phone'):
        contact_info.append(data['phone'])
    if contact_info:
        contact_para.add_run(" | ".join(contact_info)).font.size = Pt(10)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph() # Spacer

    # Skills
    if data.get('skills'):
        document.add_heading('Skills', level=2)
        skills_para = document.add_paragraph()
        skills_para.add_run(", ".join(data['skills'])).font.size = Pt(10)
        document.add_paragraph() # Spacer

    # Education
    if data.get('education'):
        document.add_heading('Education', level=2)
        for edu in data['education']:
            edu_para = document.add_paragraph()
            edu_para.add_run(f"{edu.get('degree', '')}").bold = True
            if edu.get('institution'):
                edu_para.add_run(f" - {edu.get('institution')}")
            if edu.get('passing_year'):
                edu_para.add_run(f" ({edu.get('passing_year')})")
            edu_para.runs[0].font.size = Pt(12) # Apply size to bold part
            if len(edu_para.runs) > 1: # Apply size to rest of the run
                for run in edu_para.runs[1:]:
                    run.font.size = Pt(12)

            if edu.get('marks_percentage_cgpa'):
                marks_para = document.add_paragraph()
                marks_para.add_run(f"Marks/CGPA: {edu.get('marks_percentage_cgpa')}").font.size = Pt(10)
            document.add_paragraph() # Small spacer

    # Work Experience
    if data.get('work_experience'):
        document.add_heading('Work Experience', level=2)
        for exp in data['work_experience']:
            exp_para = document.add_paragraph()
            exp_para.add_run(f"{exp.get('title', '')}").bold = True
            if exp.get('company'):
                exp_para.add_run(f" at {exp.get('company')}")
            if exp.get('dates'):
                exp_para.add_run(f" ({exp.get('dates')})")
            exp_para.runs[0].font.size = Pt(12) # Apply size to bold part
            if len(exp_para.runs) > 1: # Apply size to rest of the run
                for run in exp_para.runs[1:]:
                    run.font.size = Pt(12)

            if exp.get('description'):
                desc_lines = [line.strip() for line in exp['description'].split('\n') if line.strip()]
                for desc_line in desc_lines:
                    bullet_para = document.add_paragraph(style='List Bullet')
                    bullet_para.add_run(desc_line).font.size = Pt(10)
            document.add_paragraph() # Small spacer
            
    document.save(buffer)


if __name__ == '__main__':
    app.run(debug=True)
